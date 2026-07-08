import sys
import docx
from docx.shared import Pt
import re

md_path = r"C:\Users\kadir\.gemini\antigravity\brain\b6f34840-73d1-4d87-a8bd-2fd27b16151b\Taraftaryum_Proje_Sunum_Dokumani.md"
docx_path = r"C:\Users\kadir\Desktop\Taraftaryum_Proje_Sunum_Dokumani.docx"

doc = docx.Document()

with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

for line in lines:
    line = line.strip()
    if not line:
        continue
    
    if line.startswith('### '):
        p = doc.add_heading(line[4:], level=3)
    elif line.startswith('## '):
        p = doc.add_heading(line[3:], level=2)
    elif line.startswith('# '):
        p = doc.add_heading(line[2:], level=1)
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('|') and '---' in line:
        pass # ignore markdown table separator
    elif line.startswith('|'):
        doc.add_paragraph(line) # basic fallback for tables
    else:
        # replace basic bolding
        clean = line.replace('**', '')
        doc.add_paragraph(clean)

doc.save(docx_path)
print("Docx created at", docx_path)
