from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# T.C. University Information (Title Page)
def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("T.C.\nTEKİRDAĞ NAMIK KEMAL ÜNİVERSİTESİ\nÇORLU MÜHENDİSLİK FAKÜLTESİ\n\n\n\nBİLGİSAYAR MÜHENDİSLİĞİ PROJE-II DERSİ BİTİRME\nÇALIŞMASI\n\n\n")
    run.font.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("CANLIMAÇ: GERÇEK ZAMANLI FUTBOL VERİLERİ VE DÜNYA KUPASI UYGULAMASI\n\n")
    run.font.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run("Ad ve Soyad         : Kadir\nNo        : \n")
    
    p = doc.add_paragraph("\n\n\nDANIŞMAN Öğretim Üyesi\nDr. Öğr. Üyesi Erkan ÖZHAN\n\n\nTEKİRDAĞ-2025")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

add_title_page(doc)
add_title_page(doc) # The template has two similar title pages

# Ozet
doc.add_heading('ÖZET', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p = doc.add_paragraph()
p.add_run("Bu proje, kullanıcıların güncel futbol karşılaşmalarını, puan durumlarını ve fikstürleri gerçek zamanlı olarak takip edebilmelerini sağlayan \"CanlıMaç\" adlı web/mobil uygulamasının geliştirilmesini kapsamaktadır. Proje, frontend tarafında React Native (Expo) ve backend tarafında Python (Flask) teknolojileri kullanılarak inşa edilmiştir. Uygulama içerisinde 2026 Dünya Kupası grupları (A-L), maç sonuçları ve canlı skorlar anlık olarak gösterilmektedir. Ayrıca web scraping yöntemleriyle Transfermarkt ve ESPN gibi platformlardan veri çekilerek, İngiltere (PL), İtalya (SA), İspanya (PD), Almanya (BL1) ve Türkiye Süper Lig verileri kullanıcıya sunulmaktadır. Geliştirilen WebSocket (Socket.IO) yapısı sayesinde maç skorları ve değişiklikler anında kullanıcı arayüzüne yansıtılmakta, sayfa yenilenmeden kesintisiz bir deneyim sağlanmaktadır.")
p.add_run("\n\nAnahtar Kelimeler: React Native, Python, Flask, WebSocket, Web Scraping, Dünya Kupası")
doc.add_page_break()

# İçindekiler
doc.add_heading('İÇİNDEKİLER', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("ÖZET\n1. GİRİŞ\n2. PROJENİN TANIMI\n3. PROJENİN NEDENİ VE AMAÇLARI\n4. BULGULAR VE PROJENİN NASIL ÇALIŞTIĞI\n5. KAYNAKLAR")
doc.add_page_break()

# 1. Giris
doc.add_heading('1. GİRİŞ', level=1)
doc.add_paragraph("Günümüzde spor müsabakalarının, özellikle dünyanın en çok takip edilen sporu olan futbolun dijital ortamda anlık takibi, milyonlarca kullanıcı için vazgeçilmez bir ihtiyaç haline gelmiştir. İnternet teknolojilerinin hızla gelişmesi ve mobil cihazların yaygınlaşmasıyla birlikte, kullanıcıların bilgiye erişim beklentileri saniyeler seviyesine inmiştir. Geleneksel web siteleri ve uygulamalar, sayfa yenileme zorunluluğu veya yavaş veri güncellemeleri nedeniyle modern kullanıcının beklentilerini karşılamakta zorlanmaktadır.")
doc.add_paragraph("Bu bağlamda geliştirilen 'CanlıMaç' projesi, kullanıcılarına canlı maç skorlarını, liglerin detaylı puan durumlarını, takım kadrolarını ve güncel spor haberlerini saniyesi saniyesine aktarmayı hedefleyen yenilikçi bir uygulamadır. Sistem, verileri otomatik olarak dış kaynaklardan toplar, işler ve en düşük gecikmeyle son kullanıcıya ulaştırır. Kullanıcılar, dünyanın önde gelen Avrupa liglerinin (İngiltere, İspanya, İtalya, Almanya) yanı sıra, Türkiye Süper Ligi ve 2026 FIFA Dünya Kupası gibi büyük organizasyonların verilerini tek bir platform üzerinden detaylı, kesintisiz ve şık bir arayüzle deneyimleme imkanı bulmaktadır.")

# 2. Projenin Tanımı
doc.add_heading('2. PROJENİN TANIMI', level=1)
doc.add_paragraph("CanlıMaç, temel olarak iki ana bileşenden (frontend ve backend) oluşan, tam kapsamlı ve gerçek zamanlı bir futbol bilgi sistemidir. Uygulama, hem mobil cihazlarda (iOS, Android) hem de web tarayıcılarında çalışabilecek şekilde çapraz platform (cross-platform) yeteneklerine sahiptir.")
doc.add_paragraph("Uygulamanın İstemci (Frontend) Tarafı: İstemci tarafı, modern arayüzler oluşturmak için sektör standardı haline gelen React Native kütüphanesi ve Expo çatısı kullanılarak inşa edilmiştir. React Native, uygulamanın native (doğal) cihaz performansıyla çalışmasını sağlarken, aynı zamanda animasyonlar ve geçişlerin akıcı olmasını mümkün kılar. Kullanıcı etkileşimleri, ekranlar arası geçişler ve verilerin görselleştirilmesi bu katmanda gerçekleşir.")
doc.add_paragraph("Uygulamanın Sunucu (Backend) Tarafı: Arka planda ise Python programlama dili ile geliştirilmiş, hafif ancak güçlü bir web çatısı olan Flask yer almaktadır. Backend, tüm veri toplama, işleme, önbellekleme (caching) ve dağıtım işlemlerinin merkezidir. Dış dünyadan (ESPN, Transfermarkt, TRT Spor RSS) veri çekmek (web scraping ve API entegrasyonu) bu katmanda yapılandırılmış zamanlanmış görevler (threading) ile sağlanır.")
doc.add_paragraph("Gerçek Zamanlı İletişim (WebSocket): CanlıMaç projesini geleneksel uygulamalardan ayıran en temel teknoloji WebSocket (Socket.IO) entegrasyonudur. Backend ile frontend arasında sürekli açık kalan bu çift yönlü iletişim kanalı sayesinde, bir maçta gol olduğunda veya maç statüsü değiştiğinde sunucu, istemciye anında bir sinyal (event) gönderir. Böylece kullanıcının ekranındaki skor, sayfa yenilenmeden canlı olarak güncellenir.")

# 3. Projenin Nedeni ve Amaçları
doc.add_heading('3. PROJENİN NEDENİ VE AMAÇLARI', level=1)
doc.add_paragraph("Piyasada pek çok canlı skor uygulaması bulunmasına rağmen, bu projeyi geliştirmenin ardında yatan spesifik nedenler ve hedeflenen amaçlar şunlardır:")
doc.add_paragraph("1. Ücretsiz Servislerin Sınırlarını Aşmak (Scraping): Piyasadaki ücretsiz ve açık futbol API'leri genellikle eksik veri sunmakta veya Türkiye Süper Lig verilerini anında güncelleyememektedir. (Örneğin, Süper Lig'in 18 takımlı güncel formata geçişi birçok dış kaynakta hatalıdır). Bu projenin temel nedenlerinden biri, Transfermarkt gibi güvenilir platformlardan anlık Web Scraping (Veri Kazıma) yaparak, TFF standartlarına tam uyumlu ve %100 doğru Süper Lig puan durumu verisini kullanıcıya sunmaktır.")
doc.add_paragraph("2. Gerçek Zamanlı Kesintisiz Deneyim Sağlamak: Amaçlardan bir diğeri, HTTP isteklerinin yarattığı bekleme süresi ve sayfa yenileme zorunluluğunu ortadan kaldırmaktır. WebSocket altyapısı sayesinde kullanıcılara tamamen akıcı, skorların 'canlı' aktığı bir sistem hedeflenmiştir.")
doc.add_paragraph("3. Organizasyon Odaklı Dinamik Tasarım (2026 Dünya Kupası): Sıradan uygulamalar karmaşık lig listeleri sunarken, CanlıMaç uygulamasının amacı ana sayfasını 2026 Dünya Kupası gibi popüler turnuvalara hızlıca adapte edebilmektir. Uygulama, ana sayfasında gruplara (A'dan L'ye) doğrudan erişim imkanı vererek kullanıcıyı aradığı içeriğe en kısa yoldan ulaştırmayı hedefler.")
doc.add_paragraph("4. Performans ve Ölçeklenebilirlik: Dış kaynak API'lerine binlerce kullanıcının doğrudan istek atması sistemleri çökertebilir veya engellemelere (rate limit) sebep olabilir. Amacımız, backend'de güçlü bir önbellek (cache) sistemi kurarak, backend'in veriyi dakikada bir kendi çekmesi ve milyonlarca kullanıcıya bu önbellekten hızlıca dağıtmasıdır.")

# 4. Bulgular ve Projenin Nasil Calistigi
doc.add_heading('4. BULGULAR VE PROJENİN NASIL ÇALIŞTIĞI', level=1)
doc.add_paragraph("Bu bölümde projenin çalışma prensibi ve bir kullanıcının uygulamayı açtığı andan itibaren arka planda ve ön planda gerçekleşen olaylar, kullanılan teknolojilerle entegre biçimde adım adım açıklanmıştır:")

doc.add_heading('Uygulamayı Açılış ve İlk Veri Yüklemesi', level=2)
doc.add_paragraph("Kullanıcı CanlıMaç uygulamasını cihazında başlattığında, React Native (Expo) tabanlı arayüz saniyeler içinde yüklenir. Açılış esnasında uygulama arka planda Python Flask sunucusuna bir REST API isteği gönderir ve o günün maç programını, takımları ve mevcut durumları (skor, dakika) JSON formatında çeker. Eş zamanlı olarak Socket.IO devreye girerek sunucuyla sürekli açık kalacak bir WebSocket tüneli oluşturur.")

doc.add_heading('Ana Ekran ve Dünya Kupası Arayüzü', level=2)
doc.add_paragraph("Ana ekran açıldığında, kullanıcının karşısına özel olarak tasarlanmış 2026 Dünya Kupası Grupları (A, B, C... L) şeridi çıkar. Kullanıcı bir gruba tıkladığında:")
doc.add_paragraph("- Frontend, bu eylemi yakalar ve ilgili grubun maçlarını ekrana getirir.\n- Bu aşamada veriler daha önceden çekilip bellekte tutulduğu için geçişler bekleme ekranı olmadan anında gerçekleşir.\n- Fikstür bölümü, maçları dünün sonuçları, bugünün canlı maçları ve gelecek 3 günün karşılaşmaları olacak şekilde dinamik bir liste yapısında sunar.")

doc.add_heading('Canlı Skor Takibi (Arka Plan ve Ön Plan Uyumu)', level=2)
doc.add_paragraph("Bir maç başladığında sistem şu şekilde çalışır:")
doc.add_paragraph("1. Python Sunucusu (Backend): Her 60 saniyede bir ESPN API'sini kontrol eder (arka planda çalışan Threading mekanizması). Yeni bir gol, sarı/kırmızı kart veya dakika değişikliği varsa bunu kendi önbelleğine kaydeder.")
doc.add_paragraph("2. WebSocket Dağıtımı: Flask sunucusu, bu değişikliği fark ettiği an Socket.IO üzerinden bağlı olan tüm kullanıcılara 'match_update' sinyali gönderir.")
doc.add_paragraph("3. Kullanıcı Ekranı: Kullanıcı o an maç ekranındaysa hiçbir tuşa basmasına veya sayfayı yenilemesine gerek kalmadan ekrandaki '0-0' skoru saniyesinde '1-0' olarak güncellenir ve dakika bilgisi akmaya devam eder.")

doc.add_heading('Puan Durumu ve Detaylı Filtreleme Kullanımı', level=2)
doc.add_paragraph("Kullanıcı 'Puan Durumu' sekmesine geçtiğinde, ekranın üst kısmında lig filtreleme araçlarını (İngiltere PL, İtalya SA, İspanya PD, Almanya BL1 ve Süper Lig) görür:")
doc.add_paragraph("- Avrupa liglerine tıklandığında önbellekteki API verileri hızlıca listelenir.")
doc.add_paragraph("- Türkiye Süper Ligine tıklandığında, sistem BeautifulSoup kütüphanesi aracılığıyla geliştirilmiş Web Scraping modülünü kullanarak Transfermarkt üzerinden verileri kazır ve işler. 18 takımlı güncel puan tablosu, galibiyet, mağlubiyet, averaj gibi detaylarla birlikte ekrana mükemmel bir hizalama ile yansıtılır.")

doc.add_heading('Haber Akışı ve Diğer Modüller', level=2)
doc.add_paragraph("Uygulamanın spor haberleri sekmesi, TRT Spor gibi güvenilir kaynakların RSS beslemelerinden veri almaktadır. Backend her 6 saatte bir bu XML verisini (feedparser kütüphanesi ile) ayrıştırır, temizler ve veritabanı önbelleğine alır. Kullanıcı haber sekmesini açtığında, gereksiz reklamlar veya karmaşık bağlantılar olmadan sadece saf ve güncel spor haberlerine ulaşır.")
doc.add_paragraph("Sonuç olarak CanlıMaç uygulaması; React Native'in akıcı arayüz yetenekleri ile Python'un veri işleme (Scraping/API) gücünü WebSocket köprüsüyle kusursuz bir şekilde birleştiren, yüksek performanslı ve kullanıcı dostu modern bir yazılım ürünüdür.")

# 5. Kaynaklar
doc.add_heading('5. KAYNAKLAR', level=1)
doc.add_paragraph("[1] React Native Documentation, https://reactnative.dev/\n[2] Flask Documentation, https://flask.palletsprojects.com/\n[3] Socket.IO Documentation, https://socket.io/\n[4] Python Requests & BeautifulSoup Web Scraping Frameworks.")

doc.save('CanliMac_Proje_Raporu_v2.docx')
print("Rapor güncellendi: CanliMac_Proje_Raporu_v2.docx")
